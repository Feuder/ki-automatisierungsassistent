import logging

from pathlib import Path
from docx import Document

logger = logging.getLogger(__name__)

def datei_inhalt(dateipfad: Path) -> str:
    dateiendung = dateipfad.suffix.lower()

    if dateiendung == ".docx":
        logger.info("datei_inhalt() | DOCX erkannt | Datei: %s", dateipfad)
        dokument = Document(dateipfad)
        abschnitte: list[str] = []

        # Normale Absätze auslesen
        for absatz in dokument.paragraphs:
            text = absatz.text.strip()

            if text:
                abschnitte.append(text)

        # Tabellen auslesen
        for tabelle in dokument.tables:
            for zeile in tabelle.rows:
                zellen = [
                    zelle.text.strip()
                    for zelle in zeile.cells
                ]

                abschnitte.append("\t".join(zellen))

        text = "\n".join(abschnitte).strip()

    else:
        logger.info("datei_inhalt() | Textdatei gelesen | Datei: %s", dateipfad)
        text = dateipfad.read_text(
            encoding="utf-8",
            errors="strict",
        ).strip()

    if not text:
        logging.error(f"Es wurde kein Textinhalt gefunden: {dateipfad.name}")
        raise ValueError(
            f"Es wurde kein Textinhalt gefunden: {dateipfad.name}"
        )

    return text